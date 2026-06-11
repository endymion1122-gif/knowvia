from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "参考的文档" / "知径 Knowvia 产品定位更新说明（2026-05-30）.docx"

DEEP_INDIGO = "1B1E5F"
SOFT_VIOLET = "7B60FF"
ORBIT_BLUE = "449CFF"
PATH_TEAL = "21C7C2"
COOL_GRAY = "E6E8EF"
PALE_MINT = "E6FAF7"
PALE_LAVENDER = "EDE9FE"
SECONDARY_TEXT = "6B7280"

# compact_reference_guide preset with a named one_page_brand_brief override:
# 0.65 in margins, 9 pt body, 7.2 in table width, compact branded colors.
CONTENT_WIDTH = 7.2


def set_run_font(run, size=9, bold=False, color=DEEP_INDIGO, latin="Calibri", east_asia="PingFang SC"):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_paragraph(paragraph, before=0, after=3, line=1.05, alignment=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment
    return paragraph


def shade_cell(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill):
    properties = paragraph._p.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def mark_header_row(row):
    properties = row._tr.get_or_add_trPr()
    header = properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        properties.append(header)
    header.set(qn("w:val"), "true")


def set_cell_margin(cell, top=55, start=85, bottom=55, end=85):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = margins.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            margins.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches):
    width_dxa = int(width_inches * 1440)
    properties = cell._tc.get_or_add_tcPr()
    cell_width = properties.find(qn("w:tcW"))
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        properties.append(cell_width)
    cell_width.set(qn("w:w"), str(width_dxa))
    cell_width.set(qn("w:type"), "dxa")
    cell.width = Inches(width_inches)


def style_table(table, widths, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    table_width.set(qn("w:w"), str(round(sum(widths) * 1440)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), "85")
    table_indent.set(qn("w:type"), "dxa")
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(round(width * 1440)))
    if header:
        mark_header_row(table.rows[0])

    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[column_index])
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and row_index == 0:
                shade_cell(cell, PALE_LAVENDER)
            for paragraph in cell.paragraphs:
                set_paragraph(paragraph, after=0, line=1.0)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=8.2,
                        bold=header and row_index == 0,
                        color=DEEP_INDIGO if header and row_index == 0 else SECONDARY_TEXT,
                    )


def add_section_label(document, text):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, before=4, after=2, line=1.0)
    set_run_font(paragraph.add_run(text), size=9.4, bold=True, color=DEEP_INDIGO)


def add_callout(document, text):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, before=1, after=4, line=1.05)
    paragraph.paragraph_format.left_indent = Inches(0.10)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    shade_paragraph(paragraph, PALE_MINT)
    set_run_font(paragraph.add_run(text), size=9.2, bold=True, color=DEEP_INDIGO)


def add_inline_bullets(cell, items):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    for index, item in enumerate(items):
        if index:
            paragraph.add_run("\n")
        set_run_font(paragraph.add_run(f"• {item}"), size=8.2, color=SECONDARY_TEXT)
    set_paragraph(paragraph, after=0, line=1.0)


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    document.core_properties.title = "知径 Knowvia 产品定位更新说明"
    document.core_properties.subject = "知行星舱 StarCabin AI 知识学习产品线当前基线"
    document.core_properties.author = "StarCabin AI"
    document.core_properties.keywords = "知径 Knowvia, StarCabin AI, 产品定位, AI 阅读, 知识卡片"

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(SECONDARY_TEXT)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    set_paragraph(header, after=0, line=1.0)
    set_run_font(header.add_run("STARCABIN AI  |  PRODUCT POSITIONING UPDATE"), size=7.6, bold=True, color=SOFT_VIOLET)
    header.add_run(" " * 12)
    set_run_font(header.add_run("2026-05-30"), size=7.6, color=SECONDARY_TEXT)

    footer = section.footer.paragraphs[0]
    set_paragraph(footer, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(footer.add_run("知径 Knowvia · 让知识成为路径。"), size=7.8, color=DEEP_INDIGO)

    title = document.add_paragraph()
    set_paragraph(title, after=0, line=1.0)
    set_run_font(title.add_run("知径 "), size=24, bold=True, color=DEEP_INDIGO)
    set_run_font(title.add_run("Knowvia"), size=22, bold=True, color=DEEP_INDIGO)

    slogan = document.add_paragraph()
    set_paragraph(slogan, after=2, line=1.0)
    set_run_font(slogan.add_run("让知识成为路径。"), size=12.5, bold=True, color=DEEP_INDIGO)
    set_run_font(slogan.add_run("  Where knowledge becomes your path."), size=9.4, color=SECONDARY_TEXT)

    status = document.add_paragraph()
    set_paragraph(status, after=5, line=1.0)
    set_run_font(status.add_run("知行星舱 StarCabin AI 旗下知识学习产品线  |  "), size=8.5, color=ORBIT_BLUE)
    set_run_font(status.add_run("长期储备 · 暂缓重开发"), size=8.5, bold=True, color=PATH_TEAL)

    add_callout(
        document,
        "知径 Knowvia 是知行星舱旗下的 AI 阅读与知识卡片工具，帮助学习者将文献、网页、课程材料和笔记转化为结构化知识卡片、学习路径与可复用的知识资产。",
    )

    add_section_label(document, "01  产品矩阵中的角色")
    matrix = document.add_table(rows=4, cols=4)
    matrix_data = [
        ("产品", "产品线", "优先级", "角色"),
        ("阅衡 ReadScope", "研究型产品线", "最高", "阅读发展研究系统"),
        ("一日舱 DayCabin", "副业 MVP 产品线", "中高", "学业行动与日程管家"),
        ("知径 Knowvia", "知识学习产品线", "长期储备", "AI 阅读与知识卡片工具"),
    ]
    for row, values in zip(matrix.rows, matrix_data):
        for cell, value in zip(row.cells, values):
            cell.text = value
    style_table(matrix, [1.55, 1.65, 1.0, 3.0])

    matrix_line = document.add_paragraph()
    set_paragraph(matrix_line, before=2, after=2, line=1.0)
    set_run_font(matrix_line.add_run("矩阵表达："), size=8.5, bold=True, color=DEEP_INDIGO)
    set_run_font(matrix_line.add_run("知径管知识，一日舱管行动，阅衡管阅读发展。"), size=8.5, color=SECONDARY_TEXT)

    add_section_label(document, "02  当前边界与极简 MVP")
    scope = document.add_table(rows=4, cols=2)
    scope.cell(0, 0).text = "范围"
    scope.cell(0, 1).text = "说明"
    scope.cell(1, 0).text = "当前保留"
    scope.cell(1, 1).text = ""
    add_inline_bullets(scope.cell(1, 1), ["AI 摘要", "概念卡 / 观点卡 / 证据卡", "卡片归类", "轻量学习路径", "导出与一日舱联动概念"])
    scope.cell(2, 0).text = "当前不追求"
    scope.cell(2, 1).text = ""
    add_inline_bullets(scope.cell(2, 1), ["全能知识库平台", "Notion / Zotero / Obsidian 替代", "复杂知识图谱", "大型写作工作台", "社区与复杂分析"])
    scope.cell(3, 0).text = "未来极简 MVP"
    scope.cell(3, 1).text = "导入文本 / PDF → AI 摘要 → 生成三类知识卡片 → 保存与主题归类 → 导出 Markdown、Notion 或一日舱任务"
    style_table(scope, [1.35, 5.85])
    for index, row in enumerate(scope.rows[1:]):
        shade_cell(row.cells[0], PALE_LAVENDER if index < 2 else PALE_MINT)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=8.3, bold=True, color=DEEP_INDIGO)

    note = document.add_paragraph()
    set_paragraph(note, before=4, after=0, line=1.0)
    set_run_font(note.add_run("基线说明："), size=8.2, bold=True, color=DEEP_INDIGO)
    set_run_font(
        note.add_run("本页是 2026-05-30 起的产品定位覆盖层。原《Knowvia macOS Demo 开发说明书》继续作为已完成技术 Demo 的实现参考；发生冲突时，以本页为准。"),
        size=8.2,
        color=SECONDARY_TEXT,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
